# © 2016 cole
# License AGPL-3.0 or later (http://www.gnu.org/licenses/agpl.html).

from docxtpl import DocxTemplate

import docx

import jinja2
from datetime import datetime
from reportlab.graphics.barcode import createBarcodeDrawing

"""
使用一个独立的文件来封装需要支持图片等功能，避免污染report_docx.py
"""


def calc_length(s):
    """
    把字符串，数字类型的参数转化为docx的长度对象，如：
    12 => Pt(12)
    '12' => Pt(12)
    '12pt' => Pt(12)  单位为point
    '12cm' => Cm(12)  单位为厘米
    '12mm' => Mm(12)   单位为毫米
    '12inchs' => Inchs(12)  单位为英寸
    '12emu' => Emu(12)
    '12twips' => Twips(12)
    """
    if not isinstance(s, str):
        # 默认为像素
        return docx.shared.Pt(s)

    if s.endswith('cm'):
        return docx.shared.Cm(float(s[:-2]))
    elif s.endswith('mm'):
        return docx.shared.Mm(float(s[:-2]))
    elif s.endswith('inchs'):
        return docx.shared.Inches(float(s[:-5]))
    elif s.endswith('pt') or s.endswith('px'):
        return docx.shared.Pt(float(s[:-2]))
    elif s.endswith('emu'):
        return docx.shared.Emu(float(s[:-3]))
    elif s.endswith('twips'):
        return docx.shared.Twips(float(s[:-5]))
    else:
        # 默认为像素
        return docx.shared.Pt(float(s))


def calc_alignment(s):
    """
    把字符串转换为对齐的常量
    """
    A = docx.enum.text.WD_ALIGN_PARAGRAPH
    if s == 'center':
        return A.CENTER
    elif s == 'left':
        return A.LEFT
    elif s == 'right':
        return A.RIGHT
    else:
        return A.LEFT

@jinja2.contextfilter
def rmb_format(ctx, data):
    """
                    将数值按位数分开
    """
    value = round(data,2)
    if abs(value) < 0.01:
        # 值为0的不输出，即返回12个空格
        return ['' for i in range(12)]
    # 先将数字转为字符，去掉小数点，然后和12个空格拼成列表，取最后12个元素返回
    return (['' for i in range(12)] + list(('%0.2f' % value).replace('.', '')))[-12:]



@jinja2.contextfilter
def rmb_upper(ctx, data, field):
    """
    人民币大写
    来自：http://topic.csdn.net/u/20091129/20/b778a93d-9f8f-4829-9297-d05b08a23f80.html
    传入浮点类型的值返回 unicode 字符串
    :param 传入阿拉伯数字
    :return 返回值是对应阿拉伯数字的绝对值的中文数字
    """
    rmbmap = ["零", "壹", "贰", "叁", "肆", "伍", "陆", "柒", "捌", "玖"]
    unit = ["分", "角", "元", "拾", "佰", "仟", "万", "拾", "佰", "仟", "亿",
            "拾", "佰", "仟", "万", "拾", "佰", "仟", "兆"]
    value = round(sum(getattr(d,field) for d in data),2)
    # 冲红负数处理
    xflag = 0
    if value < 0:
        xflag = value
        value = abs(value)
    # 先把value 数字进行格式化保留两位小数，转成字符串然后去除小数点
    nums = list(map(int, list(str('%0.2f' % value).replace('.', ''))))
    words = []
    zflag = 0  # 标记连续0次数，以删除万字，或适时插入零字
    start = len(nums) - 3
    for i in range(start, -3, -1):  # 使i对应实际位数，负数为角分
        # 大部分情况对应数字不等于零 或者是刚开始循环
        if 0 != nums[start - i] or len(words) == 0:
            if zflag:
                words.append(rmbmap[0])
                zflag = 0
            words.append(rmbmap[nums[start - i]])   # 数字对应的中文字符
            words.append(unit[i + 2])               # 列表此位置的单位
        # 控制‘万/元’ 万和元比较特殊，如2拾万和2拾1万 无论有没有这个1 万字是必须的
        elif 0 == i or (0 == i % 4 and zflag < 3):
            # 上面那种情况定义了 2拾1万 的显示 这个是特殊对待的 2拾万（一类）的显示
            words.append(unit[i + 2])
            # 元（控制条件为 0 == i ）和万(控制条为(0 == i % 4 and zflag < 3))的情况的处理是一样的
            zflag = 0
        else:
            zflag += 1
    if words[-1] != unit[0]:  # 结尾非‘分’补整字 最小单位 如果最后一个字符不是最小单位(分)则要加一个整字
        words.append("整")
    if xflag < 0:             # 如果为负数则要在数字前面加上‘负’字
        words.insert(0, "负")
    return ''.join(words)

@jinja2.contextfilter
def total(ctx, data, field):
    return round(sum(getattr(d,field) for d in data),2)

@jinja2.contextfilter
def picture(ctx, data, width=None, height=None, align=None):
    """
    把图片的二进制数据（使用了base64编码）转化为一个docx.Document对象

    data：图片的二进制数据（使用了base64编码）
    word例：{{p line.goods_id.image | picture(width=’6cm’)}}
    width：图片的宽度，可以为：'12cm','12mm','12pt' 等，参考前面的 calc_length()
    height：图片的长度，如果没有设置，根据长度自动缩放
    align：图片的位置，'left'，'center'，'right'
    """

    if not data:
        return None

    # 转化为file-like对象
    # 在python2.7中，bytes==str，可以直接使用
    # 在python3.5中，bytes和str是不同的类型，需要使用base64这个库

    # data使用了base64编码，所以这里需要解码
    import base64
    data = base64.b64decode(data)

    import io
    data = io.BytesIO(data)

    tpl = ctx['tpl']
    doc = tpl.new_subdoc()

    if width:
        width = calc_length(width)
    if height:
        height = calc_length(height)

    p = doc.add_paragraph()
    p.alignment = calc_alignment(align)
    p.add_run().add_picture(data, width=width, height=height)
    return doc

@jinja2.contextfilter
def barcode(ctx, data, barcode_type, width=300, height=70, humanreadable=0, quiet=1, align=None):
    '''生成条形码、二维码
    
    在 word 中用法：
    {{p obj.name|barcode('QR',250,250)}}

    (生成 barcode 二进制后的代码，可编码成base64，再调用 picture，只是多了一次编解码过程，略慢)
    :param barcode_type: Accepted types: 'Codabar', 'Code11', 'Code128', 'EAN13', 'EAN8', 'Extended39',
    'Extended93', 'FIM', 'I2of5', 'MSI', 'POSTNET', 'QR', 'Standard39', 'Standard93',
    'UPCA', 'USPS_4State'
    :param humanreadable: Accepted values: 0 (default) or 1. 1 will insert the readable value
    at the bottom of the output image
    '''
    value = data
    if barcode_type == 'UPCA' and len(value) in (11, 12, 13):
        barcode_type = 'EAN13'
        if len(value) in (11, 12):
            value = '0%s' % value
    try:
        width, height, humanreadable, quiet = int(width), int(height), bool(int(humanreadable)), bool(int(quiet))
        # for `QR` type, `quiet` is not supported. And is simply ignored.
        # But we can use `barBorder` to get a similar behaviour.
        bar_border = 4
        if barcode_type == 'QR' and quiet:
            bar_border = 0

        barcode_img = createBarcodeDrawing(
            barcode_type, value=value, format='png', width=width, height=height,
            humanReadable=humanreadable, quiet=quiet, barBorder=bar_border
        )

        import base64
        barcode_img = base64.b64encode(barcode_img.asString('png'))

        return picture(ctx, barcode_img, width, height)
    except (ValueError, AttributeError):
        if barcode_type == 'Code128':
            raise ValueError("Cannot convert into barcode.")
        else:
            return barcode(ctx,data,'Code128', width=width, height=height,
                humanreadable=humanreadable, quiet=quiet)

def get_env():
    """
    创建一个jinja的enviroment，然后添加一个过滤器 
    """
    jinja_env = jinja2.Environment()
    jinja_env.filters['picture'] = picture
    jinja_env.filters['total'] = total
    jinja_env.filters['rmb_upper'] = rmb_upper
    jinja_env.filters['rmb_format'] = rmb_format
    jinja_env.filters['barcode'] = barcode
    jinja_env.globals['time'] = datetime.now()
    return jinja_env


def test():
    """
    演示了如何使用，可以直接执行该文件，但是需要使用自己写的docx模版，和图片
    """
    tpl = DocxTemplate("tpls/test_tpl.docx")
    # 读取图片的数据且使用base64编码
    data = open('tpls/python_logo.png', 'rb').read().encode('base64')
    obj = {'logo': data}
    # 需要添加模版对象
    ctx = {'obj': obj, 'tpl': tpl}
    jinja_env = get_env()
    tpl.render(ctx, jinja_env)

    tpl.save('tpls/test.docx')


def main():
    test()


if __name__ == '__main__':
    main()
